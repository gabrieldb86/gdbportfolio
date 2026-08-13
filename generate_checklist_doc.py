import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Manrope'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(27, 26, 25) # Graphite
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run("CHECKLIST DE ENCERRAMENTO E PUBLICAÇÃO\nPORTFÓLIO PROFISSIONAL — GABRIEL DANINO BASILIO")
    title_run.font.name = 'Manrope'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(215, 51, 50) # Carmim Red
    
    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("Guia minucioso para validação, ajustes finais, apontamento de domínio e entrada em produção.")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(90, 85, 80)
    
    doc.add_paragraph() # Spacing
    
    # Intro
    p = doc.add_paragraph(
        "Este documento foi estruturado para orientar passo a passo a verificação final de todos os elementos "
        "do portfólio de Gabriel Danino Basilio (focado em Coordenação de Conteúdo, Treinamento, Trade Marketing e T&D). "
        "Utilize este guia como um roteiro prático para garantir que cada detalhe visual, de conteúdo, de privacidade "
        "e de infraestrutura esteja perfeitamente ajustado antes da divulgação pública."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(12)
    
    # Section 1
    h1 = doc.add_heading("1. Verificação de Conteúdo e Métricas Chave", level=1)
    h1.style.font.color.rgb = RGBColor(27, 26, 25)
    
    p = doc.add_paragraph(
        "Certifique-se de que os números fundamentais e os blocos de posicionamento executivo estejam corretos "
        "e sincronizados em todas as páginas do site:"
    )
    
    items_s1 = [
        ("Métrica de Capacitação", "Validar que todas as menções exibem exatamente '114K+' (ou '114 mil+') pessoas capacitadas ao longo da carreira, sem resquícios de valores antigos como 100K ou 300K."),
        ("Anos de Experiência", "Confirmar o destaque para '17+ anos' de experiência na liderança de projetos, varejo e campo."),
        ("Operação de Campo", "Manter a referência a '130+ promotores monitorados' e 'nota média de avaliação 8,3' na seção de prova social para RH."),
        ("Posicionamento Executivo", "Validar o headline principal: 'CONTEÚDO, treinamento & TRADE MARKETING.' e a introdução baseada na execução real.")
    ]
    for title, desc in items_s1:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(f"{title}: ")
        r1.bold = True
        r2 = bp.add_run(desc)
        bp.paragraph_format.space_after = Pt(4)
        
    # Section 2
    h2 = doc.add_heading("2. Auditoria de Privacidade e Segurança de Contato", level=1)
    p = doc.add_paragraph(
        "A proteção de dados e a prevenção contra vazamentos ou golpes são premissas fundamentais deste portfólio:"
    )
    
    items_s2 = [
        ("Ausência de Telefone Visível", "Garantir que nenhum número de telefone (ex.: +55 11 94574-7353) apareça escrito como texto ou renderizado na página de CV ou de contato."),
        ("Canais Exclusivos de Contato", "Confirmar que o contato ocorre exclusivamente por e-mail profissional (gabrieldb@me.com), LinkedIn (gabrieldb86) e links diretos para o WhatsApp (que acionam o aplicativo sem expor os dígitos na interface)."),
        ("Aviso de Privacidade", "Verificar a presença do rodapé informando que os dados enviados pelo formulário são utilizados estritamente para resposta profissional.")
    ]
    for title, desc in items_s2:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(f"{title}: ")
        r1.bold = True
        r2 = bp.add_run(desc)
        bp.paragraph_format.space_after = Pt(4)

    # Section 3
    h3 = doc.add_heading("3. Responsividade e Eixo Editorial (Design Editorial)", level=1)
    p = doc.add_paragraph(
        "O portfólio utiliza uma identidade visual minimalista inspirada no editorial axis (estrita tipografia à esquerda, "
        "paleta carvão, marfim e carmim). Valide o comportamento em diferentes telas:"
    )
    
    items_s3 = [
        ("Desktop Amplo (1280px+)", "Verificar o alinhamento esquerdo consistente da barra lateral, títulos, heros e seções de projetos."),
        ("Mobile Vertical (iPhone / Android)", "Testar a abertura do menu hambúrguer, o empilhamento das 3 colunas de Focos de Coordenação e a legibilidade das fontes sem encavalamento."),
        ("Mobile Horizontal", "Garantir que a rotação da tela reorganize os blocos hero e de estatísticas sem quebrar o layout.")
    ]
    for title, desc in items_s3:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(f"{title}: ")
        r1.bold = True
        r2 = bp.add_run(desc)
        bp.paragraph_format.space_after = Pt(4)

    # Section 4
    h4 = doc.add_heading("4. Otimização para Buscadores (SEO) e ATS", level=1)
    p = doc.add_paragraph(
        "Para que recrutadores, headhunters e sistemas automatizados de triagem (ATS) encontrem o perfil com facilidade:"
    )
    
    items_s4 = [
        ("Metadados e Schema.org", "Validar que o arquivo index.html possui title, description e schema JSON-LD ProfilePage preenchidos com as palavras-chave corretas."),
        ("Palavras-Chave de Mercado", "Assegurar presença de termos como 'Instructional Design', 'ADDIE', 'Kirkpatrick', 'Trade Marketing', 'Performance de Campo' e 'T&D' no CV e na Home."),
        ("Links Externos", "Confirmar que os perfis do LinkedIn (gabrieldb86) e Behance (gabrieldb86) abrem corretamente em nova aba com rel='noreferrer'.")
    ]
    for title, desc in items_s4:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(f"{title}: ")
        r1.bold = True
        r2 = bp.add_run(desc)
        bp.paragraph_format.space_after = Pt(4)

    # Section 5
    h5 = doc.add_heading("5. Publicação, Domínio e Suporte Técnico", level=1)
    p = doc.add_paragraph(
        "Etapas finais de infraestrutura para colocar o site no ar com endereço próprio:"
    )
    
    items_s5 = [
        ("Domínio Personalizado", "Enviar ao painel do Manus os registros DNS solicitados para apontar seu domínio próprio (ex.: gabrielbasilio.com.br) para a infraestrutura de publicação."),
        ("Remoção do Script Runtime", "Caso deseje remover o script de edição ao vivo da versão estática final, abrir chamado no suporte oficial em https://help.manus.im solicitando a desativação do manus-runtime na publicação."),
        ("Monitoramento Diário", "Acompanhar os relatórios gerados pela tarefa diária configurada no Manus (às 8h de Brasília) para auditoria e tendências de mercado.")
    ]
    for title, desc in items_s5:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(f"{title}: ")
        r1.bold = True
        r2 = bp.add_run(desc)
        bp.paragraph_format.space_after = Pt(4)

    # Section 6: Action Table
    h6 = doc.add_heading("6. Tabela de Controle de Execução", level=1)
    p = doc.add_paragraph("Marque cada etapa conforme a conclusão da validação pessoal:")
    
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["Item de Verificação", "Categoria", "Status", "Notas / Ajustes"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "D73332")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p_cell = hdr_cells[i].paragraphs[0]
        p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p_cell.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            
    rows_data = [
        ("Métrica 114K+ na Home e CV", "Conteúdo", "[  ] Pendente  [X] Ok", "Validado em código e testes"),
        ("Ausência de telefone visível", "Privacidade", "[  ] Pendente  [X] Ok", "Substituído por WhatsApp/e-mail"),
        ("Responsividade Mobile", "Design/UX", "[  ] Pendente  [X] Ok", "Testado em 375px e desktop"),
        ("Metadados SEO e Schema.org", "SEO / ATS", "[  ] Pendente  [X] Ok", "Configurados em index.html"),
        ("Apontamento de Domínio Próprio", "Infraestrutura", "[  ] Pendente  [ ] Ok", "Aguardando envio de DNS"),
        ("Consulta ao suporte Manus", "Plataforma", "[  ] Pendente  [ ] Ok", "Opcional para manus-runtime")
    ]
    
    for row_idx, row_data in enumerate(rows_data):
        row_cells = table.add_row().cells
        bg = "F4EEE6" if row_idx % 2 == 1 else "FFFFFF"
        for i, text_val in enumerate(row_data):
            row_cells[i].text = text_val
            set_cell_background(row_cells[i], bg)
            set_cell_margins(row_cells[i], top=100, bottom=100, left=150, right=150)
            p_cell = row_cells[i].paragraphs[0]
            for run in p_cell.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(27, 26, 25)

    doc.add_paragraph() # Spacing
    
    # Footer note
    footer_p = doc.add_paragraph()
    f_run = footer_p.add_run("Portfólio Profissional de Gabriel Danino Basilio — Gerado para suporte à publicação em 2026.")
    f_run.font.size = Pt(9)
    f_run.font.italic = True
    f_run.font.color.rgb = RGBColor(120, 115, 110)
    
    doc.save("/home/ubuntu/gabriel-portfolio/Checklist_Publicacao_Gabriel_Basilio.docx")
    print("Document successfully created at /home/ubuntu/gabriel-portfolio/Checklist_Publicacao_Gabriel_Basilio.docx")

if __name__ == "__main__":
    create_document()
